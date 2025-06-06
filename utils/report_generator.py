from docx import Document
from io import BytesIO
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def create_docx(patient_info, diagnosis, treatment):
    """Generate a professional medical report in Word format"""
    doc = Document()
    
    # ===== Document Styles =====
    styles = doc.styles
    # Add custom styles if they don't exist
    if 'BodyText' not in styles:
        body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Calibri'
        body_style.font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(8)
    
    # ===== Title Page =====
    # Main Title
    title = doc.add_heading('COMPREHENSIVE MEDICAL REPORT', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style = title.style
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Subtitle
    doc.add_paragraph('AI-Assisted Clinical Assessment', style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()
    
    # ===== Patient Summary Section =====
    doc.add_heading('PATIENT SUMMARY', level=1)
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Light Shading Accent 1'
    
    # Fill summary table
    cells = summary_table.rows[0].cells
    cells[0].text = "Demographics"
    cells[1].text = f"{patient_info['age']} year old {patient_info['sex'].lower()}"
    
    cells = summary_table.rows[1].cells
    cells[0].text = "Vital Statistics"
    cells[1].text = f"{patient_info['height']} cm, {patient_info['weight']} kg"
    
    cells = summary_table.rows[2].cells
    cells[0].text = "Presenting Symptoms"
    cells[1].text = patient_info['symptoms']
    
    cells = summary_table.rows[3].cells
    cells[0].text = "Symptom Characteristics"
    cells[1].text = f"Duration: {patient_info['symptom_duration']}\nSeverity: {patient_info['symptom_severity']}/10"
    
    cells = summary_table.rows[4].cells
    cells[0].text = "Medical History"
    cells[1].text = patient_info['history']
    
    cells = summary_table.rows[5].cells
    cells[0].text = "Allergies"
    cells[1].text = patient_info['allergies']
    
    cells = summary_table.rows[6].cells
    cells[0].text = "Current Medications"
    cells[1].text = patient_info['medications']
    
    doc.add_paragraph()
    
    # ===== Clinical Assessment Section =====
    doc.add_heading('CLINICAL ASSESSMENT', level=1)
    doc.add_paragraph("Differential Diagnosis:", style='BodyText')
    
    # Process diagnosis content with proper formatting
    for block in diagnosis.split('\n\n'):
        if not block.strip():
            continue
            
        if block.strip().startswith('**'):
            # Section heading
            heading = doc.add_heading(level=2)
            heading.add_run(block.replace('**', '').strip()).bold = True
        elif '**' in block:
            # Bold text handling
            p = doc.add_paragraph(style='BodyText')
            parts = block.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:  # Odd indexes are between ** markers
                    run.bold = True
        elif block.strip().startswith('* '):
            # Bullet points
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(block[2:].strip())
        else:
            # Normal paragraph
            doc.add_paragraph(block, style='BodyText')
    
    doc.add_page_break()
    
    # ===== Treatment Plan Section =====
    doc.add_heading('TREATMENT PLAN', level=1)
    
    # Process treatment content
    for block in treatment.split('\n\n'):
        if not block.strip():
            continue
            
        if block.strip().startswith('**'):
            # Section heading
            heading = doc.add_heading(level=2)
            heading.add_run(block.replace('**', '').strip()).bold = True
        elif '**' in block:
            # Bold text handling
            p = doc.add_paragraph(style='BodyText')
            parts = block.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
        elif block.strip().startswith('* '):
            # Bullet points
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(block[2:].strip())
        elif ':' in block and len(block.split(':')[0]) < 30:
            # Definition-style formatting
            term, definition = block.split(':', 1)
            p = doc.add_paragraph(style='BodyText')
            p.add_run(term.strip() + ':').bold = True
            p.add_run(definition.strip())
        else:
            # Normal paragraph
            doc.add_paragraph(block, style='BodyText')
    
    # ===== Disclaimer Section =====
    doc.add_heading('IMPORTANT DISCLAIMER', level=1)
    disclaimer = doc.add_paragraph(style='BodyText')
    disclaimer.add_run("This report was generated by an AI clinical assistant and contains preliminary information only. ").bold = True
    disclaimer.add_run("The content has not been reviewed by a licensed physician and should not be considered as medical advice. Always consult with a qualified healthcare professional for diagnosis and treatment.")
    
    # Red warning box
    warning = doc.add_paragraph()
    warning.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = warning.add_run("⚠️ URGENT: Seek immediate medical attention for severe or worsening symptoms")
    run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    run.font.bold = True
    
    # ===== Final Formatting =====
    # Set consistent font for all paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if not run.font.name:
                run.font.name = 'Calibri'
            if not run.font.size:
                run.font.size = Pt(11)
    
    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio